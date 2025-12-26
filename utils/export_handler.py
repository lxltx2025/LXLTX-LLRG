"""
导出处理模块
负责将生成的综述内容导出为各种格式
"""
import os
from datetime import datetime
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from config import Config

class ExportHandler:
    def __init__(self):
        self.output_dir = Config.OUTPUT_FOLDER
        os.makedirs(self.output_dir, exist_ok=True)
    
    def export_to_markdown(self, content: str, title: str = "综述") -> str:
        """
        导出为Markdown格式
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title}_{timestamp}.md"
        filepath = os.path.join(self.output_dir, filename)
        
        # 添加元信息头
        header = f"""---
title: {title}
date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
generator: 本地综述生成系统
---

"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(header + content)
        
        return filepath
    
    def export_to_docx(self, content: str, title: str = "综述") -> str:
        """
        导出为Word文档格式
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title}_{timestamp}.docx"
        filepath = os.path.join(self.output_dir, filename)
        
        doc = Document()
        
        # 设置文档属性
        doc.core_properties.title = title
        doc.core_properties.author = "本地综述生成系统"
        
        # 添加标题
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        date_para = doc.add_paragraph(
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 空行
        
        # 解析并添加内容
        self._parse_content_to_docx(doc, content)
        
        doc.save(filepath)
        return filepath
    
    def _parse_content_to_docx(self, doc: Document, content: str):
        """
        解析Markdown格式内容并添加到Word文档
        """
        lines = content.split('\n')
        current_para = []
        
        for line in lines:
            stripped = line.strip()
            
            # 处理标题
            if stripped.startswith('# '):
                if current_para:
                    doc.add_paragraph(' '.join(current_para))
                    current_para = []
                doc.add_heading(stripped[2:], level=1)
                
            elif stripped.startswith('## '):
                if current_para:
                    doc.add_paragraph(' '.join(current_para))
                    current_para = []
                doc.add_heading(stripped[3:], level=2)
                
            elif stripped.startswith('### '):
                if current_para:
                    doc.add_paragraph(' '.join(current_para))
                    current_para = []
                doc.add_heading(stripped[4:], level=3)
                
            elif stripped.startswith('- ') or stripped.startswith('* '):
                if current_para:
                    doc.add_paragraph(' '.join(current_para))
                    current_para = []
                doc.add_paragraph(stripped[2:], style='List Bullet')
                
            elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                if current_para:
                    doc.add_paragraph(' '.join(current_para))
                    current_para = []
                # 找到点号后的内容
                idx = stripped.find('.')
                doc.add_paragraph(stripped[idx+1:].strip(), style='List Number')
                
            elif stripped == '':
                if current_para:
                    para = doc.add_paragraph(' '.join(current_para))
                    para.paragraph_format.first_line_indent = Inches(0.5)
                    current_para = []
            else:
                current_para.append(stripped)
        
        # 处理最后一段
        if current_para:
            para = doc.add_paragraph(' '.join(current_para))
            para.paragraph_format.first_line_indent = Inches(0.5)
    
    def export_to_html(self, content: str, title: str = "综述") -> str:
        """
        导出为HTML格式
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{title}_{timestamp}.html"
        filepath = os.path.join(self.output_dir, filename)
        
        # 转换Markdown为HTML
        html_content = markdown.markdown(content, extensions=['tables', 'fenced_code'])
        
        # 包装完整HTML
        html_template = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Microsoft YaHei', 'SimSun', serif;
            max-width: 800px;
            margin: 0 auto;
            padding: 40px 20px;
            line-height: 1.8;
            color: #333;
        }}
        h1 {{ text-align: center; color: #2c3e50; }}
        h2 {{ color: #34495e; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h3 {{ color: #7f8c8d; }}
        p {{ text-indent: 2em; margin: 1em 0; }}
        ul, ol {{ margin: 1em 0; padding-left: 2em; }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 1em 0;
            padding: 10px 20px;
            background: #f8f9fa;
        }}
        .meta {{
            text-align: center;
            color: #7f8c8d;
            margin-bottom: 30px;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <p class="meta">生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
    {html_content}
</body>
</html>"""
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(html_template)
        
        return filepath
    
    def export_to_pdf(self, content: str, title: str = "综述") -> Optional[str]:
        """
        导出为PDF格式（通过HTML中转）
        """
        try:
            from weasyprint import HTML
            
            # 先生成HTML
            html_path = self.export_to_html(content, title)
            
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            pdf_filename = f"{title}_{timestamp}.pdf"
            pdf_path = os.path.join(self.output_dir, pdf_filename)
            
            # 转换为PDF
            HTML(filename=html_path).write_pdf(pdf_path)
            
            # 删除临时HTML文件
            os.remove(html_path)
            
            return pdf_path
            
        except ImportError:
            print("WeasyPrint未安装，无法导出PDF")
            return None
        except Exception as e:
            print(f"PDF导出失败: {e}")
            return None
    
    def list_exports(self) -> list:
        """
        列出所有已导出的文件
        """
        exports = []
        
        for filename in os.listdir(self.output_dir):
            filepath = os.path.join(self.output_dir, filename)
            if os.path.isfile(filepath):
                stat = os.stat(filepath)
                exports.append({
                    "filename": filename,
                    "size": stat.st_size,
                    "created": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
                    "format": filename.split('.')[-1].upper()
                })
        
        return sorted(exports, key=lambda x: x['created'], reverse=True)
    
    def get_export_path(self, filename: str) -> Optional[str]:
        """
        获取导出文件的完整路径
        """
        filepath = os.path.join(self.output_dir, filename)
        if os.path.exists(filepath):
            return filepath
        return None


# 全局处理器实例
export_handler = ExportHandler()